import os
import json
import pandas as pd
from datetime import datetime
from typing import Dict, List, Any, Optional, Union
from io import BytesIO
import streamlit as st

# PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
    DocumentType = Document
except ImportError:
    DOCX_AVAILABLE = False
    DocumentType = Any

# Excel generation
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils.dataframe import dataframe_to_rows
    from openpyxl.chart import BarChart, Reference
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class ExportManager:
    """Manages export functionality for competitive analysis reports"""
    
    def __init__(self):
        self.styles = self._initialize_styles()
        self.colors = {
            'primary': '#1f77b4',
            'secondary': '#ff7f0e',
            'success': '#2ca02c',
            'warning': '#d62728',
            'info': '#17becf',
            'light': '#f8f9fa',
            'dark': '#343a40'
        }
    
    def _initialize_styles(self):
        """Initialize document styles"""
        return {
            'title': {'size': 24, 'bold': True, 'color': '#1f77b4'},
            'heading1': {'size': 18, 'bold': True, 'color': '#343a40'},
            'heading2': {'size': 16, 'bold': True, 'color': '#495057'},
            'heading3': {'size': 14, 'bold': True, 'color': '#6c757d'},
            'body': {'size': 11, 'color': '#212529'},
            'caption': {'size': 10, 'color': '#6c757d'},
            'metric': {'size': 12, 'bold': True, 'color': '#1f77b4'}
        }
    
    def get_available_formats(self) -> List[str]:
        """Get list of available export formats"""
        formats = []
        if PDF_AVAILABLE:
            formats.append('PDF')
        if DOCX_AVAILABLE:
            formats.append('Word')
        if EXCEL_AVAILABLE:
            formats.append('Excel')
        return formats
    
    def export_report(self, format_type: str, data: Dict[str, Any]) -> BytesIO:
        """Export report in specified format"""
        if format_type.upper() == 'PDF' and PDF_AVAILABLE:
            return self._export_pdf(data)
        elif format_type.upper() == 'WORD' and DOCX_AVAILABLE:
            return self._export_word(data)
        elif format_type.upper() == 'EXCEL' and EXCEL_AVAILABLE:
            return self._export_excel(data)
        else:
            raise ValueError(f"Export format '{format_type}' is not available")
    
    def _export_pdf(self, data: Dict[str, Any]) -> BytesIO:
        """Export report as PDF"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Build PDF content
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=HexColor(self.colors['primary']),
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=HexColor(self.colors['dark'])
        )
        
        # Title
        competitor_name = data.get('competitor_name', 'Unknown Competitor')
        title = Paragraph(f"Competitive Analysis Report<br/>{competitor_name}", title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        self._add_executive_summary_pdf(story, data, styles)
        story.append(PageBreak())
        
        # Business Intelligence
        if data.get('pricing_analysis') or data.get('monetization_analysis') or data.get('vision_analysis'):
            story.append(Paragraph("Business Intelligence", heading_style))
            self._add_business_intelligence_pdf(story, data, styles)
            story.append(PageBreak())
        
        # Technical Analysis
        if data.get('discovery_summary') or data.get('analyzed_pages'):
            story.append(Paragraph("Technical Analysis", heading_style))
            self._add_technical_analysis_pdf(story, data, styles)
            story.append(PageBreak())
        
        # Competitive Intelligence
        if data.get('categorization_report') or data.get('complaint_analysis'):
            story.append(Paragraph("Competitive Intelligence", heading_style))
            self._add_competitive_intelligence_pdf(story, data, styles)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _export_word(self, data: Dict[str, Any]) -> BytesIO:
        """Export report as Word document"""
        doc = Document()
        
        # Set document styles
        styles = doc.styles
        
        # Title
        competitor_name = data.get('competitor_name', 'Unknown Competitor')
        title = doc.add_heading(f'Competitive Analysis Report: {competitor_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        self._add_executive_summary_word(doc, data)
        
        # Business Intelligence
        if data.get('pricing_analysis') or data.get('monetization_analysis') or data.get('vision_analysis'):
            doc.add_page_break()
            doc.add_heading('Business Intelligence', level=1)
            self._add_business_intelligence_word(doc, data)
        
        # Technical Analysis
        if data.get('discovery_summary') or data.get('analyzed_pages'):
            doc.add_page_break()
            doc.add_heading('Technical Analysis', level=1)
            self._add_technical_analysis_word(doc, data)
        
        # Competitive Intelligence
        if data.get('categorization_report') or data.get('complaint_analysis'):
            doc.add_page_break()
            doc.add_heading('Competitive Intelligence', level=1)
            self._add_competitive_intelligence_word(doc, data)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _export_excel(self, data: Dict[str, Any]) -> BytesIO:
        """Export report as Excel workbook"""
        buffer = BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            # Overview sheet
            self._create_overview_sheet(writer, data)
            
            # Pricing Analysis sheet
            if data.get('pricing_analysis'):
                self._create_pricing_sheet(writer, data['pricing_analysis'])
            
            # Monetization Analysis sheet
            if data.get('monetization_analysis'):
                self._create_monetization_sheet(writer, data['monetization_analysis'])
            
            # Vision Analysis sheet
            if data.get('vision_analysis'):
                self._create_vision_sheet(writer, data['vision_analysis'])
            
            # Technical Analysis sheet
            if data.get('analyzed_pages'):
                self._create_technical_sheet(writer, data)
            
            # Competitive Intelligence sheet
            if data.get('categorization_report') or data.get('complaint_analysis'):
                self._create_competitive_sheet(writer, data)
        
        buffer.seek(0)
        return buffer
    
    def _add_executive_summary_pdf(self, story: List, data: Dict[str, Any], styles):
        """Add executive summary section to PDF"""
        # Basic info
        basic_info = [
            ['Competitor', data.get('competitor_name', 'Unknown')],
            ['Website', data.get('competitor_url', 'Unknown')],
            ['Country', data.get('country', 'US')],
            ['Analysis Date', datetime.now().strftime('%Y-%m-%d %H:%M')],
            ['Analysis Objective', data.get('selected_objective', 'General Analysis')]
        ]
        
        table = Table(basic_info, colWidths=[2*inch, 3*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(table)
        story.append(Spacer(1, 20))
        
        # Key metrics
        if data.get('discovery_summary'):
            story.append(Paragraph("Key Metrics", styles['Heading3']))
            metrics = [
                ['Metric', 'Value'],
                ['Total URLs Found', str(data['discovery_summary'].get('total_urls_found', 0))],
                ['Pricing Pages', str(len(data.get('discovered_urls', {}).get('pricing', [])))],
                ['Features Pages', str(len(data.get('discovered_urls', {}).get('features', [])))],
                ['Blog Pages', str(len(data.get('discovered_urls', {}).get('blog', [])))]
            ]
            
            metrics_table = Table(metrics, colWidths=[2*inch, 2*inch])
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(metrics_table)
    
    def _add_business_intelligence_pdf(self, story: List, data: Dict[str, Any], styles):
        """Add business intelligence section to PDF"""
        # Pricing Analysis
        if data.get('pricing_analysis'):
            story.append(Paragraph("Pricing Analysis", styles['Heading3']))
            pricing = data['pricing_analysis']
            
            pricing_info = [
                ['Aspect', 'Details'],
                ['Currency Detected', pricing.get('currency_detected', 'USD')],
                ['Hardware Model', pricing.get('hardware_pricing', {}).get('model_type', 'Unknown')],
                ['Software Model', pricing.get('software_pricing', {}).get('pricing_model', 'Unknown')],
                ['Hidden Fees Risk', pricing.get('hidden_fees', {}).get('risk_level', 'Unknown')]
            ]
            
            pricing_table = Table(pricing_info, colWidths=[2*inch, 3*inch])
            pricing_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(pricing_table)
            story.append(Spacer(1, 12))
        
        # Monetization Analysis
        if data.get('monetization_analysis'):
            story.append(Paragraph("Monetization Analysis", styles['Heading3']))
            monetization = data['monetization_analysis']
            
            monetization_info = [
                ['Aspect', 'Details'],
                ['Revenue Model', monetization.get('revenue_streams', {}).get('revenue_model_type', 'Unknown')],
                ['Lock-in Strength', monetization.get('lock_in_strategies', {}).get('lock_in_strength', 'Unknown')],
                ['Expansion Potential', monetization.get('expansion_revenue', {}).get('expansion_potential', 'Unknown')]
            ]
            
            monetization_table = Table(monetization_info, colWidths=[2*inch, 3*inch])
            monetization_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(monetization_table)
            story.append(Spacer(1, 12))
    
    def _add_technical_analysis_pdf(self, story: List, data: Dict[str, Any], styles):
        """Add technical analysis section to PDF"""
        if data.get('discovery_summary'):
            story.append(Paragraph("URL Discovery Results", styles['Heading3']))
            discovery = data['discovery_summary']
            
            discovery_info = [
                ['Metric', 'Value'],
                ['Total URLs Found', str(discovery.get('total_urls_found', 0))],
                ['Discovery Strategies', str(discovery.get('strategies_used', 0))],
                ['Success Rate', f"{discovery.get('success_rate', 0):.1f}%"],
                ['Total Time', f"{discovery.get('total_time', 0):.1f}s"]
            ]
            
            discovery_table = Table(discovery_info, colWidths=[2*inch, 2*inch])
            discovery_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(discovery_table)
            story.append(Spacer(1, 12))
        
        if data.get('analyzed_pages'):
            story.append(Paragraph("Content Analysis Summary", styles['Heading3']))
            pages = data['analyzed_pages']
            
            # Top 5 pages by quality
            top_pages = sorted(pages, key=lambda x: x.get('quality', {}).get('completeness_score', 0), reverse=True)[:5]
            
            page_info = [['Page Title', 'Category', 'Quality Score', 'Word Count']]
            for page in top_pages:
                page_info.append([
                    page.get('title', 'Unknown')[:30] + '...' if len(page.get('title', '')) > 30 else page.get('title', 'Unknown'),
                    page.get('category', 'Unknown'),
                    f"{page.get('quality', {}).get('completeness_score', 0):.1f}%",
                    str(page.get('word_count', 0))
                ])
            
            pages_table = Table(page_info, colWidths=[2*inch, 1*inch, 1*inch, 1*inch])
            pages_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(pages_table)
    
    def _add_competitive_intelligence_pdf(self, story: List, data: Dict[str, Any], styles):
        """Add competitive intelligence section to PDF"""
        if data.get('categorization_report'):
            story.append(Paragraph("AI Categorized Complaints", styles['Heading3']))
            categorization = data['categorization_report']
            overall_stats = categorization.get('overall_statistics', {})
            
            complaint_info = [
                ['Metric', 'Value'],
                ['Total Complaints', str(overall_stats.get('total_complaints', 0))],
                ['High/Critical Issues', str(overall_stats.get('severity_distribution', {}).get('High', 0) + overall_stats.get('severity_distribution', {}).get('Critical', 0))],
                ['High Confidence', str(overall_stats.get('confidence_distribution', {}).get('High', 0))]
            ]
            
            complaint_table = Table(complaint_info, colWidths=[2*inch, 2*inch])
            complaint_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(complaint_table)
    
    def _add_executive_summary_word(self, doc: DocumentType, data: Dict[str, Any]):
        """Add executive summary section to Word document"""
        # Basic information
        doc.add_heading('Basic Information', level=2)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        basic_info = [
            ('Competitor', data.get('competitor_name', 'Unknown')),
            ('Website', data.get('competitor_url', 'Unknown')),
            ('Country', data.get('country', 'US')),
            ('Analysis Date', datetime.now().strftime('%Y-%m-%d %H:%M')),
            ('Analysis Objective', data.get('selected_objective', 'General Analysis'))
        ]
        
        for i, (key, value) in enumerate(basic_info):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = str(value)
            
            # Make first column bold
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Key metrics
        if data.get('discovery_summary'):
            doc.add_heading('Key Metrics', level=2)
            
            metrics_table = doc.add_table(rows=5, cols=2)
            metrics_table.style = 'Table Grid'
            
            metrics = [
                ('Total URLs Found', str(data['discovery_summary'].get('total_urls_found', 0))),
                ('Pricing Pages', str(len(data.get('discovered_urls', {}).get('pricing', [])))),
                ('Features Pages', str(len(data.get('discovered_urls', {}).get('features', [])))),
                ('Blog Pages', str(len(data.get('discovered_urls', {}).get('blog', [])))),
                ('Analysis Success Rate', f"{data['discovery_summary'].get('success_rate', 0):.1f}%")
            ]
            
            for i, (key, value) in enumerate(metrics):
                row = metrics_table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = value
                
                # Make first column bold
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    def _add_business_intelligence_word(self, doc: DocumentType, data: Dict[str, Any]):
        """Add business intelligence section to Word document"""
        # Pricing Analysis
        if data.get('pricing_analysis'):
            doc.add_heading('Pricing Analysis', level=2)
            pricing = data['pricing_analysis']
            
            pricing_table = doc.add_table(rows=4, cols=2)
            pricing_table.style = 'Table Grid'
            
            pricing_info = [
                ('Currency Detected', pricing.get('currency_detected', 'USD')),
                ('Hardware Model', pricing.get('hardware_pricing', {}).get('model_type', 'Unknown')),
                ('Software Model', pricing.get('software_pricing', {}).get('pricing_model', 'Unknown')),
                ('Hidden Fees Risk', pricing.get('hidden_fees', {}).get('risk_level', 'Unknown'))
            ]
            
            for i, (key, value) in enumerate(pricing_info):
                row = pricing_table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = str(value)
                
                # Make first column bold
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Monetization Analysis
        if data.get('monetization_analysis'):
            doc.add_heading('Monetization Analysis', level=2)
            monetization = data['monetization_analysis']
            
            monetization_table = doc.add_table(rows=3, cols=2)
            monetization_table.style = 'Table Grid'
            
            monetization_info = [
                ('Revenue Model', monetization.get('revenue_streams', {}).get('revenue_model_type', 'Unknown')),
                ('Lock-in Strength', monetization.get('lock_in_strategies', {}).get('lock_in_strength', 'Unknown')),
                ('Expansion Potential', monetization.get('expansion_revenue', {}).get('expansion_potential', 'Unknown'))
            ]
            
            for i, (key, value) in enumerate(monetization_info):
                row = monetization_table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = str(value)
                
                # Make first column bold
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    def _add_technical_analysis_word(self, doc: DocumentType, data: Dict[str, Any]):
        """Add technical analysis section to Word document"""
        if data.get('discovery_summary'):
            doc.add_heading('URL Discovery Results', level=2)
            discovery = data['discovery_summary']
            
            discovery_table = doc.add_table(rows=4, cols=2)
            discovery_table.style = 'Table Grid'
            
            discovery_info = [
                ('Total URLs Found', str(discovery.get('total_urls_found', 0))),
                ('Discovery Strategies', str(discovery.get('strategies_used', 0))),
                ('Success Rate', f"{discovery.get('success_rate', 0):.1f}%"),
                ('Total Time', f"{discovery.get('total_time', 0):.1f}s")
            ]
            
            for i, (key, value) in enumerate(discovery_info):
                row = discovery_table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = value
                
                # Make first column bold
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        if data.get('analyzed_pages'):
            doc.add_heading('Content Analysis Summary', level=2)
            pages = data['analyzed_pages']
            
            # Top 5 pages by quality
            top_pages = sorted(pages, key=lambda x: x.get('quality', {}).get('completeness_score', 0), reverse=True)[:5]
            
            if top_pages:
                pages_table = doc.add_table(rows=len(top_pages) + 1, cols=4)
                pages_table.style = 'Table Grid'
                
                # Header
                header_cells = pages_table.rows[0].cells
                header_cells[0].text = 'Page Title'
                header_cells[1].text = 'Category'
                header_cells[2].text = 'Quality Score'
                header_cells[3].text = 'Word Count'
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Data rows
                for i, page in enumerate(top_pages):
                    row = pages_table.rows[i + 1]
                    row.cells[0].text = page.get('title', 'Unknown')[:30] + '...' if len(page.get('title', '')) > 30 else page.get('title', 'Unknown')
                    row.cells[1].text = page.get('category', 'Unknown')
                    row.cells[2].text = f"{page.get('quality', {}).get('completeness_score', 0):.1f}%"
                    row.cells[3].text = str(page.get('word_count', 0))
    
    def _add_competitive_intelligence_word(self, doc: DocumentType, data: Dict[str, Any]):
        """Add competitive intelligence section to Word document"""
        if data.get('categorization_report'):
            doc.add_heading('AI Categorized Complaints', level=2)
            categorization = data['categorization_report']
            overall_stats = categorization.get('overall_statistics', {})
            
            complaint_table = doc.add_table(rows=3, cols=2)
            complaint_table.style = 'Table Grid'
            
            complaint_info = [
                ('Total Complaints', str(overall_stats.get('total_complaints', 0))),
                ('High/Critical Issues', str(overall_stats.get('severity_distribution', {}).get('High', 0) + overall_stats.get('severity_distribution', {}).get('Critical', 0))),
                ('High Confidence', str(overall_stats.get('confidence_distribution', {}).get('High', 0)))
            ]
            
            for i, (key, value) in enumerate(complaint_info):
                row = complaint_table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = value
                
                # Make first column bold
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    def _create_overview_sheet(self, writer: pd.ExcelWriter, data: Dict[str, Any]):
        """Create overview sheet in Excel"""
        # Basic information
        overview_data = {
            'Attribute': ['Competitor', 'Website', 'Country', 'Analysis Date', 'Analysis Objective'],
            'Value': [
                data.get('competitor_name', 'Unknown'),
                data.get('competitor_url', 'Unknown'),
                data.get('country', 'US'),
                datetime.now().strftime('%Y-%m-%d %H:%M'),
                data.get('selected_objective', 'General Analysis')
            ]
        }
        
        df_overview = pd.DataFrame(overview_data)
        df_overview.to_excel(writer, sheet_name='Overview', index=False)
        
        # Key metrics
        if data.get('discovery_summary'):
            metrics_data = {
                'Metric': ['Total URLs Found', 'Pricing Pages', 'Features Pages', 'Blog Pages', 'Success Rate'],
                'Value': [
                    data['discovery_summary'].get('total_urls_found', 0),
                    len(data.get('discovered_urls', {}).get('pricing', [])),
                    len(data.get('discovered_urls', {}).get('features', [])),
                    len(data.get('discovered_urls', {}).get('blog', [])),
                    f"{data['discovery_summary'].get('success_rate', 0):.1f}%"
                ]
            }
            
            df_metrics = pd.DataFrame(metrics_data)
            df_metrics.to_excel(writer, sheet_name='Overview', index=False, startrow=8)
    
    def _create_pricing_sheet(self, writer: pd.ExcelWriter, pricing_analysis: Dict[str, Any]):
        """Create pricing analysis sheet in Excel"""
        # Pricing summary
        pricing_summary = {
            'Aspect': ['Currency Detected', 'Hardware Model', 'Software Model', 'Hidden Fees Risk'],
            'Details': [
                pricing_analysis.get('currency_detected', 'USD'),
                pricing_analysis.get('hardware_pricing', {}).get('model_type', 'Unknown'),
                pricing_analysis.get('software_pricing', {}).get('pricing_model', 'Unknown'),
                pricing_analysis.get('hidden_fees', {}).get('risk_level', 'Unknown')
            ]
        }
        
        df_pricing = pd.DataFrame(pricing_summary)
        df_pricing.to_excel(writer, sheet_name='Pricing Analysis', index=False)
        
        # Hidden fees details
        hidden_fees = pricing_analysis.get('hidden_fees', {}).get('fees_detected', [])
        if hidden_fees:
            fees_data = {
                'Fee Type': [fee.get('type', 'Unknown').replace('_', ' ').title() for fee in hidden_fees[:10]],
                'Description': [fee.get('description', 'No description') for fee in hidden_fees[:10]],
                'Confidence': [fee.get('confidence', 0) for fee in hidden_fees[:10]]
            }
            
            df_fees = pd.DataFrame(fees_data)
            df_fees.to_excel(writer, sheet_name='Pricing Analysis', index=False, startrow=8)
    
    def _create_monetization_sheet(self, writer: pd.ExcelWriter, monetization_analysis: Dict[str, Any]):
        """Create monetization analysis sheet in Excel"""
        # Monetization summary
        monetization_summary = {
            'Aspect': ['Revenue Model', 'Lock-in Strength', 'Expansion Potential'],
            'Details': [
                monetization_analysis.get('revenue_streams', {}).get('revenue_model_type', 'Unknown'),
                monetization_analysis.get('lock_in_strategies', {}).get('lock_in_strength', 'Unknown'),
                monetization_analysis.get('expansion_revenue', {}).get('expansion_potential', 'Unknown')
            ]
        }
        
        df_monetization = pd.DataFrame(monetization_summary)
        df_monetization.to_excel(writer, sheet_name='Monetization Analysis', index=False)
        
        # Revenue streams
        revenue_streams = monetization_analysis.get('revenue_streams', {}).get('primary_streams', [])
        if revenue_streams:
            streams_data = {
                'Stream Type': [stream.get('type', 'Unknown').replace('_', ' ').title() for stream in revenue_streams[:5]],
                'Confidence': [stream.get('confidence', 0) for stream in revenue_streams[:5]]
            }
            
            df_streams = pd.DataFrame(streams_data)
            df_streams.to_excel(writer, sheet_name='Monetization Analysis', index=False, startrow=8)
    
    def _create_vision_sheet(self, writer: pd.ExcelWriter, vision_analysis: Dict[str, Any]):
        """Create vision analysis sheet in Excel"""
        # Vision summary
        vision_summary = {
            'Aspect': ['Roadmap Signals', 'Tech Investment Areas', 'Market Expansion Signals'],
            'Count': [
                len(vision_analysis.get('product_roadmap', {}).get('upcoming_features', [])),
                len(vision_analysis.get('technology_investments', {}).get('investment_areas', [])),
                len(vision_analysis.get('market_expansion', {}).get('geographic_targets', []))
            ]
        }
        
        df_vision = pd.DataFrame(vision_summary)
        df_vision.to_excel(writer, sheet_name='Vision Analysis', index=False)
        
        # Upcoming features
        upcoming_features = vision_analysis.get('product_roadmap', {}).get('upcoming_features', [])
        if upcoming_features:
            features_data = {
                'Feature': [feature.get('feature', 'Unknown') for feature in upcoming_features[:10]],
                'Confidence': [feature.get('confidence', 0) for feature in upcoming_features[:10]]
            }
            
            df_features = pd.DataFrame(features_data)
            df_features.to_excel(writer, sheet_name='Vision Analysis', index=False, startrow=8)
    
    def _create_technical_sheet(self, writer: pd.ExcelWriter, data: Dict[str, Any]):
        """Create technical analysis sheet in Excel"""
        analyzed_pages = data.get('analyzed_pages', [])
        if analyzed_pages:
            pages_data = {
                'Page Title': [page.get('title', 'Unknown') for page in analyzed_pages],
                'Category': [page.get('category', 'Unknown') for page in analyzed_pages],
                'Quality Score': [page.get('quality', {}).get('completeness_score', 0) for page in analyzed_pages],
                'Word Count': [page.get('word_count', 0) for page in analyzed_pages],
                'URL': [page.get('url', 'Unknown') for page in analyzed_pages]
            }
            
            df_pages = pd.DataFrame(pages_data)
            df_pages.to_excel(writer, sheet_name='Technical Analysis', index=False)
    
    def _create_competitive_sheet(self, writer: pd.ExcelWriter, data: Dict[str, Any]):
        """Create competitive intelligence sheet in Excel"""
        if data.get('categorization_report'):
            categorization = data['categorization_report']
            overall_stats = categorization.get('overall_statistics', {})
            
            # Complaint summary
            complaint_summary = {
                'Metric': ['Total Complaints', 'High/Critical Issues', 'High Confidence'],
                'Value': [
                    overall_stats.get('total_complaints', 0),
                    overall_stats.get('severity_distribution', {}).get('High', 0) + overall_stats.get('severity_distribution', {}).get('Critical', 0),
                    overall_stats.get('confidence_distribution', {}).get('High', 0)
                ]
            }
            
            df_complaints = pd.DataFrame(complaint_summary)
            df_complaints.to_excel(writer, sheet_name='Competitive Intelligence', index=False)
            
            # Category distribution
            category_dist = overall_stats.get('category_distribution', {})
            if category_dist:
                cat_data = {
                    'Category': list(category_dist.keys()),
                    'Count': list(category_dist.values())
                }
                
                df_categories = pd.DataFrame(cat_data)
                df_categories.to_excel(writer, sheet_name='Competitive Intelligence', index=False, startrow=8)
    
    def get_export_filename(self, format_type: str, competitor_name: str) -> str:
        """Generate appropriate filename for export"""
        safe_name = "".join(c for c in competitor_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        date_str = datetime.now().strftime('%Y%m%d_%H%M')
        
        extensions = {
            'PDF': 'pdf',
            'WORD': 'docx',
            'EXCEL': 'xlsx'
        }
        
        ext = extensions.get(format_type.upper(), 'txt')
        return f"competitive_analysis_{safe_name}_{date_str}.{ext}" 